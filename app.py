import streamlit as st
import requests
import json
import os
import re
from docx import Document
from io import BytesIO
data_path = "latest_webhook_data.json"


# --- Template Maps ---
petition_doc_map = {
    "MVA - 1 Defendant Original Petition": "mva_1_defendant_original_petition",
    "MVA - 2 Defendants Original Petition": "mva_2_defendants_original_petition",
    "Premises Liability Original Petition": "premises_liability_original_petition",
    "Wrongful Death Original Petition": "wrongful_death_original_petition",
    "Dog Bite Original Petition": "dog_bite_original_petition",
    "Medical Malpractice Original Petition": "medical_malpractice_original_petition"
}

requests_doc_map = {
    "Plaintiff's Request for Initial Disclosures": "initial_disclosures",
    "Plaintiff's Interrogatories to Defendant": "interrogatories",
    "Plaintiff's Request for Admissions": "request_for_admissions",
    "Plaintiff's Request for Production": "request_for_production"
}

answers_doc_map = {
    "Plaintiff’s Response to Defendant’s Request for Disclosures": "answer_to_request_for_disclosures",
    "Answer to Interrogatories": "answer_to_interrogatories",
    "Answer to Request for Admissions": "answer_to_request_for_admissions",
    "Answer to Request for Production": "answer_to_request_for_production"
}

demand_letters = {
    "Stowers Demand Letter": "stowers_demand_letter",
    "General Demand Letter": "demand_letter",
    "Motor Vehicle Accident Demand Letter": "motor_vehicle_demand_letter",
    "Uninsured/Underinsured Motorist Demand Letter": "um_uim_demand_letter",
    "Slip and Fall Demand Letter": "slip_and_fall_demand_letter",
    "Dog Bite Demand Letter": "dog_bite_demand_letter"
}

insurance_docs = {
    "Letter of Representation": "letter_of_representation",
    "Uninsured/Underinsured Letter of Representation": "um_uim_letter_of_representation"
}

medical_docs = {
    "Letter of Protection": "letter_of_protection"
}

st.title("📄 Legal Document Automation")
st.divider()

# --- Document Category Selection ---
st.subheader("📂 Document Type")
selected_template_key = None
selected_doc_category = st.selectbox(
    "Choose Document Category:",
    ["Petitions", "Discovery", "Demand Letters", "Insurance", "Medical"]
)

if selected_doc_category == "Petitions":
    selected_petition_doc = st.selectbox("Select Petition Template:", list(petition_doc_map.keys()))
    selected_template_key = petition_doc_map[selected_petition_doc]
    st.divider()
elif selected_doc_category == "Discovery":
    discovery_type = st.radio("Select Discovery Task:", ["Documents to Request", "Answering Opposing Counsel Requests"])
    if discovery_type == "Documents to Request":
        selected_discovery_doc = st.selectbox("Select Document to Request:", list(requests_doc_map.keys()))
        selected_template_key = requests_doc_map[selected_discovery_doc]
    else:
        selected_discovery_doc = st.selectbox("Select Document to Answer:", list(answers_doc_map.keys()))
        selected_template_key = answers_doc_map[selected_discovery_doc]
    st.divider()
elif selected_doc_category == "Demand Letters":
    selected_demand_doc = st.selectbox("Select Demand Letter Type:", list(demand_letters.keys()))
    selected_template_key = demand_letters[selected_demand_doc]
    st.divider()
elif selected_doc_category == "Insurance":
    selected_insurance_doc = st.selectbox("Select Insurance Document:", list(insurance_docs.keys()))
    selected_template_key = insurance_docs[selected_insurance_doc]
    st.divider()
elif selected_doc_category == "Medical":
    selected_medical_doc = st.selectbox("Select Medical Document:", list(medical_docs.keys()))
    selected_template_key = medical_docs[selected_medical_doc]
    st.divider()

st.subheader("🔎 Search for Client by Name")
zapier_url = "https://hooks.zapier.com/hooks/catch/22771743/2nyirui/"

case_id    = st.text_input("Enter Case ID (optional)")
first_name = st.text_input("Client First Name")
last_name  = st.text_input("Client Last Name")

# 🌟 Quick sanity‐check to Zapier
if st.button("⚙️ Test direct to Zapier"):
    test_payload = {
        "case_id":    case_id or "TEST_CASE",
        "first_name": first_name or "TestFirst",
        "last_name":  last_name or "TestLast"
    }
    try:
        r = requests.post(zapier_url, json=test_payload, timeout=5)
        st.write("→ HTTP Status:", r.status_code)
        st.write("→ Response Text:", r.text)
    except Exception as e:
        st.error(f"Error calling Zapier: {e}")

import time

if st.button("🔍 Search Clients"):
    # 1) Send to Zapier catch hook
    search_payload = (
        {"case_id": case_id}
        if case_id
        else {"first_name": first_name, "last_name": last_name}
    )
    try:
        resp = requests.post(zapier_url, json=search_payload, timeout=5)
        resp.raise_for_status()
        st.success("✅ Search sent to Zapier; waiting for results…")
    except Exception as e:
        st.error(f"❌ Could not reach Zapier: {e}")
        st.stop()

    # 2) Wait for webhook file to be updated
    clients = []
    wait_seconds = 5
    with st.spinner(f"Waiting up to {wait_seconds}s for webhook data..."):
        for _ in range(wait_seconds):
            if os.path.exists(data_path):
                try:
                    with open(data_path, "r") as f:
                        webhook = json.load(f)
                    clients = webhook.get("clients", [])
                except Exception:
                    clients = []
            if clients:
                break
            time.sleep(1)

    # 3) Report back
    if not clients:
        st.warning("⚠️ No matching clients found in webhook_data.json")
        st.info("Try searching again, or check your Zapier Task History to confirm Zap ran.")
    else:
        st.success(f"✅ Retrieved {len(clients)} client record(s).")
        for idx, client in enumerate(clients):
            st.markdown(f"**{client.get('client_name','Unnamed')}**")
            st.markdown(f"- Accident Type: {client.get('accident_type','—')}")
            st.markdown(f"- Accident Date: {client.get('accident_date','—')}")
            if st.button("Select This Client", key=f"sel_{idx}"):
                # send the case_id back through Zapier
                try:
                    sel_resp = requests.post(zapier_url, json={"case_id": client["case_id"]})
                    sel_resp.raise_for_status()
                    st.success(f"✅ Case ID {client['case_id']} re-sent to Zapier.")
                except Exception as e:
                    st.error(f"❌ Failed to send case ID: {e}")

# --- Load Data from webhook JSON (if exists) ---
webhook_data = {}
data_path = "latest_webhook_data.json"
if os.path.exists(data_path):
    with open(data_path, "r") as f:
        try:
            webhook_data = json.load(f)
            st.success("✅ Auto-fill data loaded from webhook.")
        except json.JSONDecodeError:
            st.warning("⚠️ Could not decode JSON from webhook file.")

st.session_state["webhook_data"] = webhook_data

with st.expander("🔍 Raw Webhook Data", expanded=True):
    st.json(webhook_data)
    
    st.write("🔎 Type of webhook_data:", type(webhook_data))
    st.write("📦 Raw contents:", webhook_data)
    
replacements = {}

def get_prefill_value(key, default=""):
    data = st.session_state.get("webhook_data", {})
    
    # If it's from a "clients" array, grab the first client
    if isinstance(data, dict) and "clients" in data:
        if isinstance(data["clients"], list) and data["clients"]:
            return data["clients"][0].get(key, default)
    return data.get(key, default)

def load_template(template_name):
    path = os.path.join("templates", f"{template_name}.docx")
    if not os.path.exists(path):
        st.error(f"❌ Template not found: {template_name}.docx")
        return None
    return Document(path)

def fill_placeholders(doc, replacements):
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, val)
    return doc

PLACEHOLDER_SCHEMA = {
    "Client Info": {
        "[CLIENT_NAME]": "Client Name",
        "[CLIENT_DOB]": "Date of Birth",
        "[CLIENT_PHONE]": "Phone Number"
    },
    "Accident Info": {
        "[DATE_OF_ACCIDENT]": "Date of Accident",
        "[LOCATION_OF_ACCIDENT]": "Accident Location",
        "[POLICE_REPORT_NUMBER]": "Police Report Number"
    },
    "Attorney Info": {
        "[ATTORNEY_NAME]": "Attorney Name",
        "[FIRM_NAME]": "Firm Name"
    },
    "Insurance Info": {
        "[INSURANCE_COMPANY]": "Insurance Company",
        "[CLAIM_NUMBER]": "Claim Number"
    },
    "Legal Content": {
        "[FACTUAL_BACKGROUND]": "Factual Background",
        "[VENUE_AND_JURISDICTION]": "Venue & Jurisdiction",
        "[NEGLIGENCE_ALLEGATIONS]": "Negligence Allegations",
        "[PRAYER]": "Prayer",
        "[DAMAGES_SUMMARY]": "Damages Summary"
    },
    "Defendant Info": {
        "[DEFENDANT_1_NAME]": "Defendant 1 Name",
        "[DEFENDANT_1_ADDRESS]": "Defendant 1 Address",
        "[DEFENDANT_1_INSURANCE]": "Defendant 1 Insurance Carrier",
        "[DEFENDANT_2_NAME]": "Defendant 2 Name (if applicable)",
        "[DEFENDANT_2_ADDRESS]": "Defendant 2 Address (if applicable)",
        "[DEFENDANT_2_INSURANCE]": "Defendant 2 Insurance Carrier (if applicable)"
    }
}

GPT_SECTION_PROMPTS = {
    "[FACTUAL_BACKGROUND]": {
        "label": "Factual Background",
        "prompt": "Draft a factual background section based on the following case facts:"
    },
    "[VENUE_AND_JURISDICTION]": {
        "label": "Venue & Jurisdiction",
        "prompt": "Explain the appropriate venue and jurisdiction for this case:"
    },
    "[NEGLIGENCE_ALLEGATIONS]": {
        "label": "Negligence Allegations",
        "prompt": "List the negligence allegations against the defendant based on the following facts:"
    },
    "[PRAYER]": {
        "label": "Prayer for Relief",
        "prompt": "Draft a standard prayer for relief in a personal injury petition:"
    }
}

st.divider()
with st.expander("🧠 AI Section Generator (Factual Background, Venue, Negligence, Prayer)"):
    if "gpt_sections" not in st.session_state:
        st.session_state["gpt_sections"] = {}

    for placeholder, meta in GPT_SECTION_PROMPTS.items():
        st.markdown(f"### 📄 {meta['label']}")
        context = st.text_area(f"Enter context for {meta['label']}:", key=f"ctx_{placeholder}")

        if st.button(f"Generate {meta['label']}", key=f"btn_{placeholder}"):
            result = f"[Generated GPT Section for {meta['label']}\n\n{context}]"

            st.session_state["gpt_sections"][placeholder] = result

        if placeholder in st.session_state["gpt_sections"]:
            st.text_area(
                f"🧠 Generated {meta['label']} Output",
                st.session_state["gpt_sections"][placeholder],
                height=200,
                key=f"out_{placeholder}"
            )
            replacements[placeholder] = st.session_state["gpt_sections"][placeholder]

st.divider()
with st.expander("📍 Venue & Jurisdiction Generator (optional override)"):
    venue_zip = st.text_input("Enter ZIP code of the accident location")
    defendant_county = st.text_input("Enter county where Defendant resides (if known)")
    defendant_principal_office = st.text_input("Enter county of Defendant's principal office (if applicable)")

    def generate_venue_narrative(accident_county, def_county=None, office_county=None):
        venue_bases = []
        if accident_county:
            venue_bases.append(f"under CPRC §15.002(a)(1) because a substantial part of the events giving rise to this lawsuit occurred in {accident_county} County")
        if def_county:
            venue_bases.append(f"under CPRC §15.002(a)(2) because the Defendant resides in {def_county} County")
        if office_county:
            venue_bases.append(f"under CPRC §15.002(a)(3) because the Defendant’s principal office is located in {office_county} County")
        return f"Venue is proper in {accident_county} County, Texas, and also potentially " + "; ".join(venue_bases) + "."

    if st.button("Generate Venue Narrative"):
        accident_county = "Harris" if venue_zip == "77002" else "Unknown"
        venue_narrative = generate_venue_narrative(
            accident_county,
            defendant_county,
            defendant_principal_office
        )
        st.text_area("Generated Venue & Jurisdiction", venue_narrative, height=200)
        replacements["[VENUE_AND_JURISDICTION]"] = venue_narrative

st.divider()
st.subheader("📝 Input Client Information Manually")
for section, fields in PLACEHOLDER_SCHEMA.items():
    with st.expander(f"📂 {section}"):
        show_extra = True
        if section == "Defendant Info":
            show_extra = st.checkbox("Include Second Defendant?", key="show_def2")
        for placeholder, label in fields.items():
            if "DEFENDANT_2" in placeholder and not st.session_state.get("show_def2", False):
                continue
            default_val = get_prefill_value(label.lower().replace(" ", "_"))
            value = st.text_input(label, value=default_val, key=placeholder)
            replacements[placeholder] = value

if selected_template_key:
    doc = load_template(selected_template_key)
    if doc:
        filled_doc = fill_placeholders(doc, replacements)
        if st.button("📄 Preview Document Text"):
            preview = "\n".join([p.text for p in filled_doc.paragraphs])

            st.text_area("Document Preview", preview, height=400)

        buffer = BytesIO()
        filled_doc.save(buffer)
        st.download_button(
            label="📥 Download Final Document",
            data=buffer.getvalue(),
            file_name=f"{selected_template_key}_final.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


































